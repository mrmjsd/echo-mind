import os
import fitz  # PyMuPDF
from docx import Document
from pathlib import Path
import nltk
from nltk.tokenize import sent_tokenize
from langchain.schema import Document as LangDocument
from core.config import DOCS_DIR
from core.logging_config import get_logger

# Setup logger
logger = get_logger(__name__)

# Download tokenizer once (safe if already downloaded)
try:
    nltk.download('punkt_tab', quiet=True)
except Exception:
    # Non-fatal: will fail at tokenize if missing
    logger.warning("Failed to download NLTK tokenizer 'punkt_tab'. Proceeding anyway.")

def get_available_files():
    docs_dir = DOCS_DIR
    if not os.path.exists(docs_dir):
        return []
    supported_extensions = ('.pdf', '.txt', '.doc', '.docx')
    return [f for f in os.listdir(docs_dir) if Path(f).suffix.lower() in supported_extensions]

def extract_text(filename, content=None):
    docs_dir = DOCS_DIR
    text = ""

    if content is not None:
        # Handle content passed directly (e.g., upload)
        temp_path = f"temp_{filename}"
        with open(temp_path, "wb") as f:
            f.write(content)
        file_path = temp_path
    else:
        # Load from local folder
        file_path = os.path.join(docs_dir, filename)
        if not os.path.exists(file_path):
            return []

    file_ext = Path(filename).suffix.lower()

    try:
        if file_ext == '.pdf':
            doc = fitz.open(file_path)
            text = "\n".join([page.get_text() for page in doc])
            doc.close()
        elif file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        elif file_ext in ('.doc', '.docx'):
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        else:
            text = ""
    except Exception as e:
        logger.exception("Error reading document: %s", filename)
        text = ""

    if content is not None:
        try:
            os.remove(file_path)
        except Exception:
            logger.warning("Failed to remove temporary file: %s", file_path)

    # Normalize and split into sentences
    try:
        raw_sentences = sent_tokenize(text.replace('\n', ' '))
    except LookupError:
        logger.error("NLTK tokenizer not available. Ensure 'punkt' or appropriate model is installed.")
        raw_sentences = [text]
    sentences = [s.strip() for s in raw_sentences if len(s.strip()) > 20]

    # Chunking: Group 2–3 sentences together
    chunks = []
    chunk = []
    max_chunk_length = 500  # Approx. character limit per chunk

    for sentence in sentences:
        chunk.append(sentence)
        if sum(len(s) for s in chunk) > max_chunk_length or len(chunk) >= 3:
            chunks.append(" ".join(chunk).strip())
            chunk = []

    if chunk:
        chunks.append(" ".join(chunk).strip())

    # Return LangChain Document format with metadata
    return [
        LangDocument(page_content=chunk, metadata={"source": filename})
        for chunk in chunks if len(chunk.strip()) > 20
    ]